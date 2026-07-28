from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT = r'c:\xampp\htdocs\Luminest\Property_Manager_Module_Explainer.docx'

def add_code_label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.size = Pt(10)
    return p


def add_code_explanation(doc, bullets):
    for bullet in bullets:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bullet)


def build_document():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Property Manager Module Code Explanation')
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('Luminest').italic = True

    doc.add_paragraph(
        'This document explains the Property Manager modules by code block, focusing on what each block does, how AJAX is used, and how data flows between the browser and the database.'
    )

    # Dashboard
    doc.add_heading('1. dashboard.php', level=1)
    add_code_label(doc, 'Code block: session and access guard')
    add_code_explanation(doc, [
        'Starts the session, loads the database connection, and blocks users who are not Property_Manager.',
        'This ensures only the correct role can view the dashboard.'
    ])
    add_code_label(doc, 'Code block: helper functions')
    add_code_explanation(doc, [
        'tableExists() checks whether a table exists before any query runs.',
        'getColumns() reads the table schema so the page can adapt to missing or renamed columns.',
        'This makes the dashboard resilient if the database evolves.'
    ])
    add_code_label(doc, 'Code block: getDashboardData()')
    add_code_explanation(doc, [
        'Builds the dashboard metrics for houses and maintenance requests.',
        'Counts house statuses such as available, reserved, and sold.',
        'Counts maintenance requests by status and also fetches the five most recent requests for the preview list.',
        'Uses dynamic column detection so it can work with either request_id or id, and with different maintenance schema variants.'
    ])
    add_code_label(doc, 'Code block: AJAX dashboard_data endpoint')
    add_code_explanation(doc, [
        'When the page is called with ?ajax=dashboard_data, the PHP side returns JSON instead of HTML.',
        'The front end uses this response to refresh counts without reloading the page.'
    ])
    add_code_label(doc, 'Code block: HTML dashboard cards and recent maintenance list')
    add_code_explanation(doc, [
        'Renders the metric cards for total, available, reserved, sold, pending, in-progress, and completed counts.',
        'Shows a short recent-maintenance list with status badges.'
    ])
    add_code_label(doc, 'Code block: JavaScript refresh logic')
    add_code_explanation(doc, [
        'Calls dashboard.php?ajax=dashboard_data on button click and every 30 seconds.',
        'Updates the visible counters and the recent request list from JSON.',
        'This is the main AJAX-driven refresh loop on the dashboard.'
    ])

    # Maintenance active module
    doc.add_heading('2. maintenance.php', level=1)
    add_code_label(doc, 'Code block: session guard and schema helpers')
    add_code_explanation(doc, [
        'Starts the session and blocks non-Property_Manager users.',
        'Defines helpers for table existence, column lookup, and schema-aware expression building.',
        'These helpers let the page work even if the maintenance table uses different column names.'
    ])
    add_code_label(doc, 'Code block: getStaffList()')
    add_code_explanation(doc, [
        'Fetches all Maintenance_Staff users with their expertise.',
        'The expertise field is used so the manager can assign a staff member who matches the request category.'
    ])
    add_code_label(doc, 'Code block: fetchMaintenanceRequests()')
    add_code_explanation(doc, [
        'Builds the active maintenance query with search, status, and priority filters.',
        'Joins tenant and staff names so the table can show human-readable values.',
        'Excludes completed requests so this page only shows active work.'
    ])
    add_code_label(doc, 'Code block: POST action=update_request AJAX handler')
    add_code_explanation(doc, [
        'Accepts a single row update from the browser and returns JSON.',
        'Validates the request ID, current record, status, and priority.',
        'Validates assigned staff, checks role = Maintenance_Staff, and compares staff expertise with the request category.',
        'Auto-moves a pending request to accepted when it gets assigned for the first time.',
        'Updates the database and returns a success message for the UI.'
    ])
    add_code_label(doc, 'Code block: GET ajax=search handler')
    add_code_explanation(doc, [
        'Returns filtered maintenance rows as JSON for live search and filter changes.',
        'This is what powers the search box and dropdown filters without a page reload.'
    ])
    add_code_label(doc, 'Code block: active maintenance table markup')
    add_code_explanation(doc, [
        'Displays request title, tenant, status, priority, assigned staff, and update time.',
        'Each row includes editable dropdowns for status, priority, and staff assignment.'
    ])
    add_code_label(doc, 'Code block: JavaScript table refresh and change handler')
    add_code_explanation(doc, [
        'searchInput, statusFilter, and priorityFilter all trigger loadData().',
        'loadData() calls the AJAX search endpoint and rebuilds the table body from JSON.',
        'The change listener sends a FormData POST to update_request when a row field changes.',
        'After success, the page refreshes the table so the user immediately sees the updated state.'
    ])

    # Maintenance history
    doc.add_heading('3. maintenance_history.php', level=1)
    add_code_label(doc, 'Code block: session guard and schema helpers')
    add_code_explanation(doc, [
        'Same role protection and schema-safe helper pattern as the active maintenance page.',
        'This keeps the history page stable even if the database schema varies.'
    ])
    add_code_label(doc, 'Code block: fetchMaintenanceHistory()')
    add_code_explanation(doc, [
        'Queries only resolved and completed requests.',
        'Joins tenant and staff names so history rows remain readable.',
        'Supports search by title, category, tenant, and staff, plus a finished-status filter.',
        'Orders results by completion or resolution time so the newest finished items appear first.'
    ])
    add_code_label(doc, 'Code block: AJAX search handler')
    add_code_explanation(doc, [
        'Returns history rows as JSON for live filtering.',
        'The browser can search the archive without reloading the page.'
    ])
    add_code_label(doc, 'Code block: history table and JavaScript')
    add_code_explanation(doc, [
        'The table presents historical requests with resolved and completed timestamps.',
        'The JavaScript search input and status dropdown trigger AJAX refreshes.',
        'This page acts as the finished-work archive for the Property Manager.'
    ])

    # Maintenance staff page
    doc.add_heading('4. maintenance_staff.php', level=1)
    add_code_label(doc, 'Code block: session guard and schema helpers')
    add_code_explanation(doc, [
        'Protects the page so only Property_Manager can access it.',
        'Detects which maintenance assignment and request ID columns exist in the current database.'
    ])
    add_code_label(doc, 'Code block: fetchStaff()')
    add_code_explanation(doc, [
        'Loads all Maintenance_Staff users.',
        'Counts active requests separately from completed jobs using the request status column.',
        'This gives the manager a live view of workload and finished work per staff member.'
    ])
    add_code_label(doc, 'Code block: remove_staff_role AJAX handler')
    add_code_explanation(doc, [
        'Accepts a staff user ID from the browser and changes the role back to Prospect.',
        'Returns JSON so the page can update immediately after removal.'
    ])
    add_code_label(doc, 'Code block: search AJAX handler and table UI')
    add_code_explanation(doc, [
        'Search requests are handled via ?ajax=search and returned as JSON.',
        'The front-end table is rebuilt from the JSON response and shows Active Assignments and Completed Jobs.',
        'A history shortcut links directly to the maintenance archive page.'
    ])

    # Reservation list
    doc.add_heading('5. reservation_list.php', level=1)
    add_code_label(doc, 'Code block: POST AJAX update_status handler')
    add_code_explanation(doc, [
        'Receives reservation status updates from the modal form using AJAX.',
        'Validates the reservation ID and allowed status values.',
        'Uses a transaction so reservation, house status, and user role stay in sync.',
        'When accepted or paid, the house becomes reserved or sold and the user becomes Tenant.'
    ])
    add_code_label(doc, 'Code block: reservation query and status badge helper')
    add_code_explanation(doc, [
        'Fetches reservations with guest and house details for the table.',
        'getStatusBadgeHtml() maps reservation status values to Bootstrap badges.'
    ])
    add_code_label(doc, 'Code block: table, modal, and AJAX form script')
    add_code_explanation(doc, [
        'Each row has an Update button that opens a modal for status changes.',
        'The modal form submits asynchronously and updates the badge in place.',
        'This is one of the main examples of AJAX-driven inline management in the Property Manager area.'
    ])

    # Tenants
    doc.add_heading('6. tenants.php', level=1)
    add_code_label(doc, 'Code block: AJAX search and initial tenant query')
    add_code_explanation(doc, [
        'Builds a tenant list by joining users to house_reservations.',
        'The ?ajax_search=1 path returns JSON search results for the live filter box.',
        'Search covers tenant identity, contact data, and reservation details.'
    ])
    add_code_label(doc, 'Code block: table rendering and modal details')
    add_code_explanation(doc, [
        'The table shows tenant info, house type, block/lot, reservation status, and reserved date.',
        'A Details button opens a modal with the selected tenant data.',
        'The JavaScript redraws the table from AJAX results and updates the visible count.'
    ])

    # Listings
    doc.add_heading('7. listings.php', level=1)
    add_code_label(doc, 'Code block: fetchListings()')
    add_code_explanation(doc, [
        'Returns house inventory rows with owner details.',
        'Supports search and status filtering for house type, block, lot, owner, and inventory state.'
    ])
    add_code_label(doc, 'Code block: POST update_status handler')
    add_code_explanation(doc, [
        'Processes house status updates through AJAX.',
        'Accepts only available, reserved, and sold.',
        'Returns a JSON success message so the UI can update without reload.'
    ])
    add_code_label(doc, 'Code block: status form table and JavaScript')
    add_code_explanation(doc, [
        'Each listing row has a small inline form for changing the house status.',
        'The JavaScript search and status filter rebuild the table from AJAX JSON results.',
        'Submitting the inline form updates the house row in place and refreshes the filtered list.'
    ])

    doc.add_heading('8. Overall Property Manager Flow', level=1)
    add_code_explanation(doc, [
        'Dashboard summarizes the system with periodic AJAX refreshes.',
        'Reservations, tenants, listings, maintenance, and staff pages all use asynchronous requests for search or update actions.',
        'The maintenance workflow is the most stateful: tenant submits, property manager assigns, staff resolves, tenant completes, and the history page archives the finished request.',
        'Most Property Manager pages follow the same pattern: validate session, read schema safely, run a JSON endpoint, and refresh the table on the client side.'
    ])

    doc.save(OUTPUT)


if __name__ == '__main__':
    build_document()
